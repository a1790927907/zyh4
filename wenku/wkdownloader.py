#coding=utf-8
import requests
import re,json,os
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.132 Safari/537.36"
}
def getdocwithtxt(url):
    urlcookie = 'https://www.baidu.com/'
    res = requests.get(urlcookie, headers=headers)
    cookies = res.cookies
    try:
        url = url.replace('%3A', ':').replace('%2F', "/")
        res = requests.get(url,headers=headers,cookies=cookies)
    except:
        return ('error', 'none')
    try:
        urllist = re.findall(r'pageLoadUrl\\("|x22):\\("|x22)(.*?)\\("|x22)}',res.text)
    except:
        return ('error', 'none')
    try:
        pagenum = re.findall(r"\'totalPageNum\'.*?\'(.*?)\'",res.text)[0]
    except:
        pagenum = 0
    try:
        title = re.findall(r'<title>(.*?)</title>', res.text)[0]
        title = title.replace('&nbsp;', '').replace('- 百度文库', '')
    except:
        title = '未知标题'
    targeturllist = []
    for urltuple in urllist:
        for target in urltuple:
            if re.search(r'^http',target) != None:
                targeturllist.append(str(target).replace("\\",""))
    trueurllist = []
    for url2 in targeturllist:
        if "0.json" in url2:
            trueurllist.append(url2)
    if len(trueurllist) == int(pagenum)*2:
        trueurllist = trueurllist[:int(len(trueurllist)/2)]
    result = ""
    for fomaturl in trueurllist:
        resformat = requests.get(fomaturl,headers=headers).text
        jsonstr = re.findall(r'wenku_[0-9]+\((.*)\)',resformat)[0]
        jsondata = json.loads(jsonstr)
        count = 0
        for text in jsondata["body"]:
            try:
                if text["c"] == " ":
                    result += "\n\n"
                else:
                    result += text["c"]
                '''
                if text["c"] == " " and count == 0:
                    result += "\n"
                    count += 1
                elif text["c"] == " ":
                    result += text["c"]
                    count = 0
                else:
                    result += text["c"]
                '''
            except:
                continue
    result = result.replace('百度文库','').replace('让每个人平等地提升自我','')
    with open(f'./static_all/file/{title}.txt',"w",encoding="utf-8") as f:
        f.write(result)
    return ('success',title)
def getdocwithword(url):
    urlcookie = 'https://www.baidu.com/'
    res = requests.get(urlcookie, headers=headers)
    cookies = res.cookies
    try:
        url = url.replace('%3A', ':').replace('%2F', "/")
        res = requests.get(url,headers=headers,cookies=cookies)
    except:
        return ('error', 'none')
    try:
        urllist = re.findall(r'pageLoadUrl\\("|x22):\\("|x22)(.*?)\\("|x22)}',res.text)
    except:
        return ('error', 'none')
    try:
        pagenum = re.findall(r"\'totalPageNum\'.*?\'(.*?)\'",res.text)[0]
    except:
        pagenum = 0
    try:
        title = re.findall(r'<title>(.*?)</title>', res.text)[0]
        title = title.replace('&nbsp;', '').replace('- 百度文库', '')
    except:
        title = '未知标题'
    targeturllist = []
    for urltuple in urllist:
        for target in urltuple:
            if re.search(r'^http',target) != None:
                targeturllist.append(str(target).replace("\\",""))
    trueurllist = []
    for url2 in targeturllist:
        if "0.json" in url2:
            trueurllist.append(url2)
    if len(trueurllist) == int(pagenum)*2:
        trueurllist = trueurllist[:int(len(trueurllist)/2)]
    result = ""
    for fomaturl in trueurllist:
        resformat = requests.get(fomaturl,headers=headers).text
        jsonstr = re.findall(r'wenku_[0-9]+\((.*)\)',resformat)[0]
        jsondata = json.loads(jsonstr)
        count = 0
        for text in jsondata["body"]:
            try:
                if text["c"] == " ":
                    result += "\n\n"
                else:
                    result += text["c"]
                '''
                if text["c"] == " " and count == 0:
                    result += "\n"
                    count += 1
                elif text["c"] == " ":
                    result += text["c"]
                    count = 0
                else:
                    result += text["c"]
                '''
            except:
                continue
    result = result.replace('百度文库', '').replace('让每个人平等地提升自我', '').replace('\n\n\n\n\n','')
    document = Document()
    # 加入不同等级的标题
    head = document.add_heading('', 1)
    head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head = head.add_run(f'{title}')
    head.font.name = '黑体'
    h = head._element
    h.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 添加文本
    paragraph = document.add_paragraph('')
    paragraph.paragraph_format.line_spacing = 1.5
    # 设置字号
    run = paragraph.add_run(result)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 保存文件
    document.save(f'./static_all/file/{title}.docx')
    return ('success',title)

if __name__ == '__main__':
    while True:
        print('PPT类文档暂时不支持...')
        url = input("输入文库URL：")
        proto = input('1.txt版,2.word版：')
        if proto == '1':
            getdocwithtxt(url)
        elif proto == '2':
            print(getdocwithword(url))
        else:
            getdocwithword(url)







'''
https://wkbjcloudbos.bdimg.com/v1/docconvert9342/wk/9ba6d785aff8886c0311c806eeff9c47/0.json?
responseContentType=application%2Fjavascript
&responseCacheControl=max-age%3D3888000
&responseExpires=Thu%2C%2002%20Jul%202020%2001%3A13%3A20%20%2B0800
&authorization=bce-auth-v1%2Ffa1126e91489401fa7cc85045ce7179e%2F2020-05-17T17%3A13%3A20Z%2F3600%2Fhost%2Ffb460f10e6538e6d0e1015359ddbb6a181c6b79cbdd23021c5495ffa317c2612
&x-bce-range=0-9533
&token=eyJ0eXAiOiJKSVQiLCJ2ZXIiOiIxLjAiLCJhbGciOiJIUzI1NiIsImV4cCI6MTU4OTczOTIwMCwidXJpIjp0cnVlLCJwYXJhbXMiOlsicmVzcG9uc2VDb250ZW50VHlwZSIsInJlc3BvbnNlQ2FjaGVDb250cm9sIiwicmVzcG9uc2VFeHBpcmVzIiwieC1iY2UtcmFuZ2UiXX0%3D.wUlDMWi%2Bxqjkir9i3o73Fnbd2sOpyn6LPy18%2FTu0qUg%3D.1589739200

https://wkbjcloudbos.bdimg.com/v1/docconvert521/wk/d6113009a56fcc4b6211bd56bf483def/0.json?responseContentType=application%2Fjavascript&responseCacheControl=max-age%3D3888000&responseExpires=Thu%2C%2002%20Jul%202020%2001%3A14%3A17%20%2B0800&authorization=bce-auth-v1%2Ffa1126e91489401fa7cc85045ce7179e%2F2020-05-17T17%3A14%3A17Z%2F3600%2Fhost%2Fb63da668eba9d8310f7ecedf80cc8a8b6f5520629b7453b6616a31c73ce26ce8&x-bce-range=0-11806&token=eyJ0eXAiOiJKSVQiLCJ2ZXIiOiIxLjAiLCJhbGciOiJIUzI1NiIsImV4cCI6MTU4OTczOTI1NywidXJpIjp0cnVlLCJwYXJhbXMiOlsicmVzcG9uc2VDb250ZW50VHlwZSIsInJlc3BvbnNlQ2FjaGVDb250cm9sIiwicmVzcG9uc2VFeHBpcmVzIiwieC1iY2UtcmFuZ2UiXX0%3D.kHzVaYnuzagjH%2BXpjD7y%2BKfGnBQsEXoA8qAS79mEoPg%3D.1589739257
'''